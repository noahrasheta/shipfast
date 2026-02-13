"""Tests for the Word document converter."""

from pathlib import Path

import docx
import pytest

from converters import ConfidenceLevel, ExtractionResult
from converters.word import (
    WordConverter,
    _clean_cell,
    _count_images,
    _heading_level,
    _paragraph_to_markdown,
    _score_confidence,
    _table_to_markdown,
)


# ------------------------------------------------------------------
# Fixtures
# ------------------------------------------------------------------


@pytest.fixture
def converter():
    return WordConverter()


@pytest.fixture
def simple_docx(tmp_path):
    """Create a simple .docx with a heading and a few paragraphs."""
    doc = docx.Document()
    doc.add_heading("Site Overview", level=1)
    doc.add_paragraph("This is a data center opportunity located in Dallas, TX.")
    doc.add_paragraph("The facility has 50MW of available power capacity.")
    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_headings(tmp_path):
    """Create a .docx with multiple heading levels."""
    doc = docx.Document()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("This document summarizes the opportunity.")
    doc.add_heading("Location Details", level=2)
    doc.add_paragraph("Pioneer Park, Dallas, TX 75201.")
    doc.add_heading("Power Infrastructure", level=2)
    doc.add_paragraph("50MW from Oncor with planned expansion to 130MW.")
    doc.add_heading("Substation Details", level=3)
    doc.add_paragraph("Dedicated 138kV substation on-site.")
    path = tmp_path / "headings.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_table(tmp_path):
    """Create a .docx with a table."""
    doc = docx.Document()
    doc.add_heading("Specifications", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Attribute"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Total Power"
    table.cell(1, 1).text = "50 MW"
    table.cell(2, 0).text = "Floor Space"
    table.cell(2, 1).text = "120,000 sqft"
    table.cell(3, 0).text = "Cooling"
    table.cell(3, 1).text = "Chilled Water"
    doc.add_paragraph("Additional notes about the specifications.")
    path = tmp_path / "with_table.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_multiple_tables(tmp_path):
    """Create a .docx with text and multiple tables interleaved."""
    doc = docx.Document()
    doc.add_heading("Due Diligence Report", level=1)
    doc.add_paragraph("Summary of key findings.")

    doc.add_heading("Financial Summary", level=2)
    table1 = doc.add_table(rows=3, cols=3)
    table1.cell(0, 0).text = "Year"
    table1.cell(0, 1).text = "Revenue"
    table1.cell(0, 2).text = "EBITDA"
    table1.cell(1, 0).text = "2024"
    table1.cell(1, 1).text = "$5M"
    table1.cell(1, 2).text = "$1.2M"
    table1.cell(2, 0).text = "2025"
    table1.cell(2, 1).text = "$7.5M"
    table1.cell(2, 2).text = "$2.1M"

    doc.add_heading("Site Metrics", level=2)
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "Metric"
    table2.cell(0, 1).text = "Value"
    table2.cell(1, 0).text = "PUE"
    table2.cell(1, 1).text = "1.25"

    doc.add_paragraph("End of report.")
    path = tmp_path / "multi_table.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx(tmp_path):
    """Create an empty .docx with no content."""
    doc = docx.Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_image(tmp_path):
    """Create a .docx with an embedded image."""
    from PIL import Image as PILImage
    import io

    # Create a small test image.
    img = PILImage.new("RGB", (100, 100), color="red")
    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)

    doc = docx.Document()
    doc.add_heading("Site Photos", level=1)
    doc.add_paragraph("Below are photos of the facility.")
    doc.add_picture(img_bytes, width=docx.shared.Inches(2))
    doc.add_paragraph("Photo shows the exterior of the building.")
    path = tmp_path / "with_image.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_special_chars(tmp_path):
    """Create a .docx with pipe characters and other markdown-sensitive chars."""
    doc = docx.Document()
    doc.add_heading("Special Characters", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Test|Item"
    table.cell(1, 1).text = "Line1\nLine2"
    path = tmp_path / "special.docx"
    doc.save(str(path))
    return path


# ------------------------------------------------------------------
# Unit tests for helper functions
# ------------------------------------------------------------------


class TestHeadingLevel:
    def test_heading_1(self):
        assert _heading_level("Heading 1") == 1

    def test_heading_2(self):
        assert _heading_level("Heading 2") == 2

    def test_heading_3(self):
        assert _heading_level("Heading 3") == 3

    def test_heading_6(self):
        assert _heading_level("Heading 6") == 6

    def test_heading_beyond_6_clamped(self):
        assert _heading_level("Heading 9") == 6

    def test_heading_0_clamped_to_1(self):
        assert _heading_level("Heading 0") == 1

    def test_no_number_defaults_to_1(self):
        assert _heading_level("Heading") == 1

    def test_custom_heading_name(self):
        assert _heading_level("Custom Heading") == 1


class TestCleanCell:
    def test_strips_whitespace(self):
        assert _clean_cell("  padded  ") == "padded"

    def test_replaces_newlines(self):
        assert _clean_cell("line1\nline2") == "line1 line2"

    def test_replaces_carriage_return(self):
        assert _clean_cell("line1\rline2") == "line1 line2"

    def test_empty_string(self):
        assert _clean_cell("") == ""

    def test_normal_text(self):
        assert _clean_cell("Hello World") == "Hello World"


class TestScoreConfidence:
    def test_no_content(self):
        confidence, reason = _score_confidence(0, 0, 0, 0)
        assert confidence == ConfidenceLevel.LOW
        assert "no content" in reason

    def test_very_little_content(self):
        confidence, reason = _score_confidence(10, 1, 0, 0)
        assert confidence == ConfidenceLevel.LOW
        assert "very little" in reason

    def test_sparse_content(self):
        confidence, reason = _score_confidence(50, 3, 0, 0)
        assert confidence == ConfidenceLevel.MEDIUM
        assert "sparse" in reason

    def test_substantial_content(self):
        confidence, reason = _score_confidence(500, 10, 2, 3)
        assert confidence == ConfidenceLevel.HIGH
        assert "well-structured" in reason

    def test_high_with_tables(self):
        confidence, reason = _score_confidence(200, 5, 3, 1)
        assert confidence == ConfidenceLevel.HIGH
        assert "table" in reason

    def test_high_with_headings(self):
        confidence, reason = _score_confidence(200, 5, 0, 4)
        assert confidence == ConfidenceLevel.HIGH
        assert "heading" in reason


# ------------------------------------------------------------------
# Converter behaviour tests
# ------------------------------------------------------------------


class TestWordConverterBasics:
    def test_supported_extensions(self, converter):
        assert ".docx" in converter.supported_extensions

    def test_can_handle_docx(self, converter):
        assert converter.can_handle(Path("report.docx")) is True
        assert converter.can_handle(Path("REPORT.DOCX")) is True

    def test_cannot_handle_non_word(self, converter):
        assert converter.can_handle(Path("data.pdf")) is False
        assert converter.can_handle(Path("data.xlsx")) is False
        assert converter.can_handle(Path("data.pptx")) is False
        assert converter.can_handle(Path("data.doc")) is False

    def test_missing_file_returns_error(self, converter):
        result = converter.convert(Path("/tmp/nonexistent_word_12345.docx"))
        assert result.success is False
        assert result.confidence == ConfidenceLevel.LOW
        assert "not found" in result.error.lower()
        assert result.method == "python-docx"

    def test_wrong_extension_returns_error(self, converter, tmp_path):
        txt = tmp_path / "test.txt"
        txt.write_text("not a word doc")
        result = converter.convert(txt)
        assert result.success is False
        assert "not a word" in result.error.lower()

    def test_result_type(self, converter):
        result = converter.convert(Path("/tmp/nonexistent_12345.docx"))
        assert isinstance(result, ExtractionResult)

    def test_source_path_resolved(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        assert result.source_path == simple_docx.resolve()


class TestDocxExtraction:
    def test_simple_file(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        assert result.success is True
        assert result.method == "python-docx"
        assert result.confidence in (ConfidenceLevel.HIGH, ConfidenceLevel.MEDIUM)

    def test_contains_heading(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        assert "# Site Overview" in result.text

    def test_contains_paragraphs(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        assert "data center opportunity" in result.text
        assert "50MW" in result.text

    def test_multiple_heading_levels(self, converter, docx_with_headings):
        result = converter.convert(docx_with_headings)
        assert result.success is True
        assert "# Executive Summary" in result.text
        assert "## Location Details" in result.text
        assert "## Power Infrastructure" in result.text
        assert "### Substation Details" in result.text

    def test_heading_content_preserved(self, converter, docx_with_headings):
        result = converter.convert(docx_with_headings)
        assert "Pioneer Park" in result.text
        assert "50MW from Oncor" in result.text
        assert "138kV substation" in result.text

    def test_table_extraction(self, converter, docx_with_table):
        result = converter.convert(docx_with_table)
        assert result.success is True
        assert "Attribute" in result.text
        assert "Value" in result.text
        assert "50 MW" in result.text
        assert "120,000 sqft" in result.text
        assert "Chilled Water" in result.text

    def test_table_is_markdown_formatted(self, converter, docx_with_table):
        result = converter.convert(docx_with_table)
        assert "|" in result.text
        assert "---" in result.text

    def test_multiple_tables(self, converter, docx_with_multiple_tables):
        result = converter.convert(docx_with_multiple_tables)
        assert result.success is True
        assert "Revenue" in result.text
        assert "EBITDA" in result.text
        assert "PUE" in result.text
        assert "1.25" in result.text
        assert result.metadata["table_count"] == 2

    def test_metadata(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        assert "paragraph_count" in result.metadata
        assert "table_count" in result.metadata
        assert "heading_count" in result.metadata
        assert "image_count" in result.metadata
        assert "total_chars" in result.metadata

    def test_heading_count_in_metadata(self, converter, docx_with_headings):
        result = converter.convert(docx_with_headings)
        assert result.metadata["heading_count"] >= 4

    def test_table_count_in_metadata(self, converter, docx_with_table):
        result = converter.convert(docx_with_table)
        assert result.metadata["table_count"] == 1

    def test_confidence_high_for_rich_document(self, converter, docx_with_headings):
        result = converter.convert(docx_with_headings)
        assert result.confidence == ConfidenceLevel.HIGH


class TestDocxImages:
    def test_image_noted(self, converter, docx_with_image):
        result = converter.convert(docx_with_image)
        assert result.success is True
        assert "embedded image" in result.text.lower()
        assert result.metadata["image_count"] >= 1

    def test_image_not_extracted_as_text(self, converter, docx_with_image):
        result = converter.convert(docx_with_image)
        assert "not extracted as text" in result.text.lower()

    def test_no_image_note_when_none(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        assert "embedded image" not in result.text.lower()
        assert result.metadata["image_count"] == 0


class TestDocxEdgeCases:
    def test_empty_document(self, converter, empty_docx):
        result = converter.convert(empty_docx)
        assert result.success is True
        assert result.confidence == ConfidenceLevel.LOW

    def test_special_characters_in_table(self, converter, docx_with_special_chars):
        result = converter.convert(docx_with_special_chars)
        assert result.success is True
        # Pipe should be escaped in the markdown table.
        assert "Test\\|Item" in result.text
        # Newline in cell should be replaced with space.
        assert "Line1 Line2" in result.text

    def test_corrupted_file_handled(self, converter, tmp_path):
        """A corrupted file should return an error result, not crash."""
        bad_file = tmp_path / "corrupted.docx"
        bad_file.write_bytes(b"this is not a real docx file")
        result = converter.convert(bad_file)
        assert result.success is False
        assert result.confidence == ConfidenceLevel.LOW
        assert result.error is not None
        assert "extraction failed" in result.error.lower()

    def test_confidence_summary_works(self, converter, simple_docx):
        result = converter.convert(simple_docx)
        summary = result.confidence_summary
        assert "python-docx" in summary
        assert "confidence" in summary


class TestDocxImportFromPackage:
    def test_import_from_converters(self):
        from converters import WordConverter
        converter = WordConverter()
        assert ".docx" in converter.supported_extensions
